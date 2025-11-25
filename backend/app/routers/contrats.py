# from fastapi import APIRouter, Depends, HTTPException
# from sqlalchemy.orm import Session
# from app.db import get_db
# from app.models.models import Contrat, Employee
# from app.schemas import ContratCreate, ContratOut

# router = APIRouter(tags=["Contrats"])

# # ============================
# #  Créer un contrat
# # ============================
# @router.post("/", response_model=ContratOut)
# def create_contrat(contrat: ContratCreate, db: Session = Depends(get_db)):
#     """Créer un nouveau contrat pour un employé existant"""
#     employee = db.query(Employee).filter(Employee.id == contrat.employee_id).first()
#     if not employee:
#         raise HTTPException(status_code=404, detail="Employé non trouvé")

#     new_contrat = Contrat(
#         employee_id=contrat.employee_id,
#         type_contrat=contrat.type_contrat,
#         date_debut=contrat.date_debut,
#         date_fin=contrat.date_fin,
#         salaire=contrat.salaire,
#     )
#     db.add(new_contrat)
#     db.commit()
#     db.refresh(new_contrat)
#     return new_contrat


# # ============================
# #  Lister tous les contrats avec nom complet et poste
# # ============================
# @router.get("/")
# def get_contrats(db: Session = Depends(get_db)):
#     """Récupérer tous les contrats avec le nom complet et le poste de l'employé"""
#     contrats = (
#         db.query(Contrat, Employee.fullname, Employee.poste)
#         .join(Employee, Contrat.employee_id == Employee.id)
#         .all()
#     )

#     result = []
#     for contrat, fullname, poste in contrats:
#         result.append({
#             "id": contrat.id,
#             "employee_id": contrat.employee_id,
#             "type_contrat": contrat.type_contrat,
#             "date_debut": contrat.date_debut,
#             "date_fin": contrat.date_fin,
#             "salaire": contrat.salaire,
#             "nom_complet": fullname if fullname and fullname.strip() != "" else "Inconnu",
#             "poste": poste if poste and poste.strip() != "" else "Non défini"
#         })
#     return result



















# from fastapi import APIRouter, Depends, HTTPException
# from sqlalchemy.orm import Session
# from datetime import date
# from app.db import get_db
# from app.models.models import Contrat, Employee
# from app.schemas.contrats import ContratCreate, ContratOut

# router = APIRouter(tags=["Contrats"])

# # ---------------------------
# # Calcul automatique préavis / indemnités
# # ---------------------------
# def calculate_preavis_indemnites(type_contrat: str):
#     preavis = ""
#     indemnites = ""

#     if type_contrat == "CDI":
#         preavis = "30 jours"
#         indemnites = ""  # CDI n'a pas d'indemnités obligatoires
#     elif type_contrat == "CDD":
#         preavis = "15 jours avant échéance"
#         indemnites = "10% de la rémunération brute totale"
#     elif type_contrat == "Stage":
#         preavis = "48 heures"
#         indemnites = ""
#     elif type_contrat == "Alternance":
#         preavis = "1 semaine"
#         indemnites = ""

#     return preavis, indemnites

# # ============================
# # Créer un contrat
# # ============================
# @router.post("/", response_model=ContratOut)
# def create_contrat(contrat: ContratCreate, db: Session = Depends(get_db)):
#     employee = db.query(Employee).filter(Employee.id == contrat.employee_id).first()
#     if not employee:
#         raise HTTPException(status_code=404, detail="Employé non trouvé")

#     # Calcul automatique
#     preavis_auto, indemnites_auto = calculate_preavis_indemnites(contrat.type_contrat)

#     new_contrat = Contrat(
#         employee_id=contrat.employee_id,
#         type_contrat=contrat.type_contrat,
#         date_debut=contrat.date_debut,
#         date_fin=contrat.date_fin,
#         salaire=contrat.salaire,
#         poste=contrat.poste,
#         periode=contrat.periode,
#         avantages=contrat.avantages,
#         clauses=contrat.clauses,
#         type_travail=contrat.type_travail or "Temps plein",
#         preavis=preavis_auto,
#         indemnites=indemnites_auto
#     )

#     db.add(new_contrat)
#     db.commit()
#     db.refresh(new_contrat)

#     return new_contrat

# # ============================
# # Lister tous les contrats
# # ============================
# @router.get("/", response_model=list[ContratOut])
# def get_contrats(db: Session = Depends(get_db)):
#     contrats = db.query(Contrat).all()
#     result = []

#     for c in contrats:
#         employee = db.query(Employee).filter(Employee.id == c.employee_id).first()
#         nom_complet = f"{employee.nom} {employee.prenom}" if employee else "Inconnu"

#         result.append(
#             ContratOut(
#                 id=c.id,
#                 employee_id=c.employee_id,
#                 type_contrat=c.type_contrat,
#                 date_debut=c.date_debut,
#                 date_fin=c.date_fin,
#                 salaire=c.salaire,
#                 poste=c.poste,
#                 periode=c.periode,
#                 avantages=c.avantages,
#                 clauses=c.clauses,
#                 type_travail=c.type_travail,
#                 preavis=c.preavis,
#                 indemnites=c.indemnites
#             )
#         )
#     return result














# from fastapi import APIRouter, Depends, HTTPException
# from fastapi.responses import StreamingResponse
# from sqlalchemy.orm import Session
# from datetime import date
# from io import BytesIO
# from docx import Document
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import A4
# import os

# from app.db import get_db
# from app.models.models import Contrat, Employee
# from app.schemas.contrats import ContratCreate, ContratOut

# router = APIRouter(tags=["Contrats"])

# # ---------------------------
# # Calcul automatique préavis / indemnités
# # ---------------------------
# def calculate_preavis_indemnites(type_contrat: str):
#     preavis = ""
#     indemnites = ""

#     if type_contrat == "CDI":
#         preavis = "30 jours"
#         indemnites = ""
#     elif type_contrat == "CDD":
#         preavis = "15 jours avant échéance"
#         indemnites = "10% de la rémunération brute totale"
#     elif type_contrat == "Stage":
#         preavis = "48 heures"
#         indemnites = ""
#     elif type_contrat == "Alternance":
#         preavis = "1 semaine"
#         indemnites = ""

#     return preavis, indemnites

# # ============================
# # Créer un contrat
# # ============================
# @router.post("/", response_model=ContratOut)
# def create_contrat(contrat: ContratCreate, db: Session = Depends(get_db)):
#     employee = db.query(Employee).filter(Employee.id == contrat.employee_id).first()
#     if not employee:
#         raise HTTPException(status_code=404, detail="Employé non trouvé")

#     preavis_auto, indemnites_auto = calculate_preavis_indemnites(contrat.type_contrat)

#     new_contrat = Contrat(
#         employee_id=contrat.employee_id,
#         type_contrat=contrat.type_contrat,
#         date_debut=contrat.date_debut,
#         date_fin=contrat.date_fin,
#         salaire=contrat.salaire,
#         poste=contrat.poste,
#         periode=contrat.periode,
#         avantages=contrat.avantages,
#         clauses=contrat.clauses,
#         type_travail=contrat.type_travail or "Temps plein",
#         preavis=preavis_auto,
#         indemnites=indemnites_auto
#     )

#     db.add(new_contrat)
#     db.commit()
#     db.refresh(new_contrat)

#     return new_contrat

# # ============================
# # Lister tous les contrats
# # ============================
# @router.get("/", response_model=list[ContratOut])
# def get_contrats(db: Session = Depends(get_db)):
#     contrats = db.query(Contrat).all()
#     result = []

#     for c in contrats:
#         employee = db.query(Employee).filter(Employee.id == c.employee_id).first()
#         nom_complet = f"{employee.nom} {employee.prenom}" if employee else "Inconnu"

#         result.append(
#             ContratOut(
#                 id=c.id,
#                 employee_id=c.employee_id,
#                 type_contrat=c.type_contrat,
#                 date_debut=c.date_debut,
#                 date_fin=c.date_fin,
#                 salaire=c.salaire,
#                 poste=c.poste,
#                 periode=c.periode,
#                 avantages=c.avantages,
#                 clauses=c.clauses,
#                 type_travail=c.type_travail,
#                 preavis=c.preavis,
#                 indemnites=c.indemnites
#             )
#         )
#     return result

# # ============================
# # Export dynamique PDF / DOCX
# # ============================
# EXPORT_DIR = "archives/contrats"
# os.makedirs(EXPORT_DIR, exist_ok=True)

# @router.get("/export")
# def export_contrat(employee_id: int, format: str = "pdf", db: Session = Depends(get_db)):
#     employee = db.query(Employee).filter(Employee.id == employee_id).first()
#     if not employee:
#         raise HTTPException(status_code=404, detail="Employé non trouvé")

#     contrat = db.query(Contrat).filter(Contrat.employee_id == employee_id).order_by(Contrat.id.desc()).first()
#     if not contrat:
#         raise HTTPException(status_code=404, detail="Contrat non trouvé")

#     # Texte dynamique
#     texte = f"""
#     CONTRAT DE TRAVAIL

#     Nom: {employee.nom} {employee.prenom}
#     Poste: {contrat.poste}
#     Salaire: {contrat.salaire} €
#     Date de début: {contrat.date_debut.strftime('%d/%m/%Y') if contrat.date_debut else '—'}
#     Type de contrat: {contrat.type_contrat}
#     Période: {contrat.periode}
#     Avantages: {contrat.avantages}
#     Clauses: {contrat.clauses}
#     Préavis: {contrat.preavis}
#     Indemnités: {contrat.indemnites}
#     """

#     filename = f"{EXPORT_DIR}/Contrat_{employee.nom}_{employee.prenom}_{date.today()}.{format.lower()}"

#     if format.lower() == "pdf":
#         buffer = BytesIO()
#         c = canvas.Canvas(buffer, pagesize=A4)
#         for i, line in enumerate(texte.split("\n")):
#             c.drawString(50, 800 - i * 20, line)
#         c.showPage()
#         c.save()
#         buffer.seek(0)

#         # Archivage
#         with open(filename, "wb") as f:
#             f.write(buffer.getbuffer())

#         return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=Contrat_{employee.nom}_{employee.prenom}.pdf"})

#     elif format.lower() == "docx":
#         doc = Document()
#         for line in texte.split("\n"):
#             doc.add_paragraph(line)
#         doc.save(filename)

#         # Envoyer le fichier
#         file_stream = open(filename, "rb")
#         return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                                  headers={"Content-Disposition": f"attachment; filename=Contrat_{employee.nom}_{employee.prenom}.docx"})

#     else:
#         raise HTTPException(status_code=400, detail="Format invalide. Choisir 'pdf' ou 'docx'.")







from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.db import get_db
from app.models.models import Contrat, Employee
from app.schemas.contrats import ContratCreate, ContratOut
import tempfile
from docx import Document
from fpdf import FPDF
import os

router = APIRouter(tags=["Contrats"])

# ---------------------------
# Calcul automatique préavis / indemnités
# ---------------------------
def calculate_preavis_indemnites(type_contrat: str):
    preavis = ""
    indemnites = ""

    if type_contrat == "CDI":
        preavis = "30 jours"
    elif type_contrat == "CDD":
        preavis = "15 jours avant échéance"
        indemnites = "10% de la rémunération brute totale"
    elif type_contrat == "Stage":
        preavis = "48 heures"
    elif type_contrat == "Alternance":
        preavis = "1 semaine"

    return preavis, indemnites

# ============================
# Créer un contrat
# ============================
@router.post("/", response_model=ContratOut)
def create_contrat(contrat: ContratCreate, db: Session = Depends(get_db)):
    employee = db.query(Employee).filter(Employee.id == contrat.employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Employé non trouvé")

    preavis_auto, indemnites_auto = calculate_preavis_indemnites(contrat.type_contrat)

    new_contrat = Contrat(
        employee_id=contrat.employee_id,
        type_contrat=contrat.type_contrat,
        date_debut=contrat.date_debut,
        date_fin=contrat.date_fin,
        salaire=contrat.salaire,
        poste=contrat.poste,
        periode=contrat.periode,
        avantages=contrat.avantages,
        clauses=contrat.clauses,
        type_travail=contrat.type_travail or "Temps plein",
        preavis=preavis_auto,
        indemnites=indemnites_auto
    )

    db.add(new_contrat)
    db.commit()
    db.refresh(new_contrat)

    return new_contrat

# ============================
# Lister tous les contrats
# ============================
@router.get("/", response_model=list[ContratOut])
def get_contrats(db: Session = Depends(get_db)):
    contrats = db.query(Contrat).all()
    result = []

    for c in contrats:
        employee = db.query(Employee).filter(Employee.id == c.employee_id).first()
        nom_complet = f"{employee.nom} {employee.prenom}" if employee else "Inconnu"

        result.append(
            ContratOut(
                id=c.id,
                employee_id=c.employee_id,
                type_contrat=c.type_contrat,
                date_debut=c.date_debut,
                date_fin=c.date_fin,
                salaire=c.salaire,
                poste=c.poste,
                periode=c.periode,
                avantages=c.avantages,
                clauses=c.clauses,
                type_travail=c.type_travail,
                preavis=c.preavis,
                indemnites=c.indemnites
            )
        )
    return result

# ============================
# Export contrat PDF/DOCX
# ============================
@router.get("/export/{contrat_id}")
def export_contrat(contrat_id: int, format: str = "pdf", db: Session = Depends(get_db)):
    contrat = db.query(Contrat).filter(Contrat.id == contrat_id).first()
    if not contrat:
        raise HTTPException(status_code=404, detail="Contrat non trouvé")

    employee = db.query(Employee).filter(Employee.id == contrat.employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Employé non trouvé")

    # Texte dynamique
    texte = f"""
    Contrat de travail

    Employé: {employee.nom} {employee.prenom}
    Poste: {contrat.poste}
    Salaire: {contrat.salaire}
    Date début: {contrat.date_debut}
    Type: {contrat.type_contrat}
    Préavis: {contrat.preavis}
    Indemnités: {contrat.indemnites}
    """

    # Fichier temporaire
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{format}")

    try:
        if format.lower() == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in texte.split("\n"):
                pdf.cell(0, 10, txt=line, ln=True)
            pdf.output(temp_file.name)
        elif format.lower() == "docx":
            doc = Document()
            for line in texte.split("\n"):
                doc.add_paragraph(line)
            doc.save(temp_file.name)
        else:
            raise HTTPException(status_code=400, detail="Format non supporté")

        return FileResponse(path=temp_file.name, filename=f"contrat_{contrat_id}.{format}", media_type="application/octet-stream")
    finally:
        # Les fichiers temporaires seront supprimés automatiquement après envoi
        pass
